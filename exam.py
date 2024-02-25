from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns


class Exam():
    def __init__(self, config, questions, exam_number, solutions) -> None:
        # Importazione della configurazione.
        self.config = config

        # Importazione delle domande assegnate a questo esame.
        self.questions = questions

        # Importazione della numerazione assegnata a questo esame.
        self.exam_number = exam_number

        # Importazione della scelta di generare il correttore.
        self.solutions = solutions

        # Creazione e settaggio di un nuovo documento Word per l'esame.
        self.doc = Document()
        self.set_document(self.doc, self.config['title'], self.config['heading'])

        # Creazione e settaggio di un nuovo documento Word per il correttore.
        if self.solutions:
            self.doc_solutions = Document()
            self.set_document(self.doc_solutions, self.config['title'], self.config['heading'])
        

    # Metodo per il settaggio preliminare del documento Word.
    def set_document(self, doc, title, heading) -> None:
        if self.config["number_heading"]:
            par = doc.add_paragraph(title + f" - #{self.exam_number}")
        else:
            par = doc.add_paragraph(title)
        par.style = doc.styles['Title']
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = heading        
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
        doc.add_section(WD_SECTION.CONTINUOUS)
        section = doc.sections[1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')


    # Metodo per il settaggio della numerazione automatica delle pagine.
    def add_page_number(self, run) -> None:
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    
    def write(self) -> None:
        # Vengono scritte le singole domande all'interno del documento e nel relativo correttore.
        for index, question in enumerate(self.questions):
            # Se la configurazione lo consente, aggiunge il numero ad ogni domanda.
            if self.config["number_questions"]:
                self.doc.add_heading(f"{index + 1}) {question['question']}", 3)
                if self.solutions:
                    self.doc_solutions.add_heading(f"{index + 1}) {question['question']}", 3)
            
            # Altrimenti, viene scritto solo il testo della domanda.
            else:
                self.doc.add_heading(f"{question['question']}", 3)
                if self.solutions:
                    self.doc_solutions.add_heading(f"{question['question']}", 3)
            
            # Vengono scritte le opzioni di risposta della domanda.
            for i in range(0, self.config['options_supported']):
                self.doc.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
                if self.solutions:
                    runner = self.doc_solutions.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
                    runner.bold = question['options'][i]['correct']
        
            self.doc.save(f"{self.config['destination_path']}/{self.config['file_name']}_{str(self.exam_number)}.docx")
            if self.solutions:
                self.doc_solutions.save(f"{self.config['destination_path']}/{self.config['file_name']}_{str(self.exam_number)}_solutions.docx")
            