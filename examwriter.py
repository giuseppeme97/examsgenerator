from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor
import docx


class ExamWriter():
    def __init__(self, config, questions, exam_number) -> None:
        self.config = config
        self.questions = questions
        self.exam_number = exam_number
        self.doc = Document()
        self.set_document(self.doc, self.config['title'], self.config['heading'], self.config["number_heading"])

        if self.config["solutions"]:
            self.doc_solutions = Document()
            self.set_document(self.doc_solutions, self.config['title'], self.config['heading'], self.config["number_heading"])

        self.write()
        

    def set_document(self, doc, title, heading, number_heading) -> None:
        styles_element = doc.styles.element
        rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
        lang_default = rpr_default.xpath('w:lang')[0]
        lang_default.set(docx.oxml.shared.qn('w:val'),'it-IT')
        
        par = doc.add_paragraph(title + f" - #{self.exam_number}") if number_heading else doc.add_paragraph(title)
 
        #par.style = doc.styles['Title']
        for run in par.runs:
            run.bold = True
            run.font.size = Pt(15)
        
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = doc.sections[0]
        upper = section.header
        paragraph = upper.paragraphs[0]
        paragraph.text = heading        
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
        doc.add_section(WD_SECTION.CONTINUOUS)
        section = doc.sections[1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')


    def add_page_number(self, doc) -> None:
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        doc._r.append(fldChar1)
        doc._r.append(instrText)
        doc._r.append(fldChar2)


    def sanitize(self, h):
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
            run.font.size = Pt(11)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    
    def write(self) -> None:
        for index, question in enumerate(self.questions):
            heading = (f"{index + 1}) " if self.config["number_questions"] else "") + question['question']
            h = self.doc.add_heading(heading, 3)
            self.sanitize(h)
            if self.config["solutions"]:
                h = self.doc_solutions.add_heading(heading, 3)
                self.sanitize(h)
            
            for i in range(0, self.config['options_supported']):
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run(question['options'][i]['text'])
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                if self.config["solutions"]:
                    p = self.doc_solutions.add_paragraph(style='List Bullet')
                    r = p.add_run(question['options'][i]['text'])
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    r.bold = question['options'][i]['correct']
                    r.underline = question['options'][i]['correct']
        
            self.doc.save(f"{self.config['destination_path']}/{self.config['file_name']}_{str(self.exam_number)}.docx")
            if self.config["solutions"]:
                self.doc_solutions.save(f"{self.config['destination_path']}/{self.config['file_name']}_{str(self.exam_number)}_solutions.docx")
            